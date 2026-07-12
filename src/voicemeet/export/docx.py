"""DOCX export — session to .docx file via python-docx."""

from __future__ import annotations

from pathlib import Path

from voicemeet.export import _resolve_output
from voicemeet.store.models import Segment, Session
from voicemeet.summarize.ollama_summarizer import SummaryResult


def export_docx(
    session: Session,
    summary: SummaryResult,
    segments: list[Segment],
    output_path: str | Path | None = None,
    export_dir: Path | None = None,
) -> str:
    """Export session as a DOCX file. Returns the file path."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX export. Install with: pip install python-docx"
        ) from e

    path = _resolve_output(output_path, session, summary, "docx", export_dir)

    doc = Document()

    # Title
    title = summary.title or session.title or "Untitled Meeting"
    doc.add_heading(title, level=1)

    # Metadata
    started = session.started_dt
    date_str = started.strftime("%d.%m.%Y")
    start_time = started.strftime("%H:%M")
    end_time = session.ended_dt.strftime("%H:%M") if session.ended_dt else "—"

    meta_lines = [
        f"Date: {date_str}",
        f"Time: {start_time} – {end_time} ({session.duration_str})",
        f"Mode: {session.mode}",
    ]
    participants = summary.participants or session.participants
    if participants:
        meta_lines.append(f"Participants: {', '.join(participants)}")
    topics = summary.topics or session.topics
    if topics:
        meta_lines.append(f"Topics: {', '.join(topics)}")

    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(11)

    doc.add_paragraph()

    # Summary
    if summary.summary_markdown:
        doc.add_heading("Summary", level=2)
        for line in summary.summary_markdown.split("\n"):
            line = line.strip()
            if line:
                clean = line.lstrip("#").strip()
                if clean:
                    doc.add_paragraph(clean)

    # Action Items
    if summary.action_items:
        doc.add_heading("Action Items", level=2)
        for item in summary.action_items:
            doc.add_paragraph(item, style="List Bullet")

    # Transcript
    if segments:
        doc.add_heading("Transcript", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Time"
        hdr[1].text = "Speaker"
        hdr[2].text = "Text"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for seg in segments:
            row = table.add_row().cells
            row[0].text = seg.start_str
            row[1].text = seg.speaker or "Speaker"
            row[2].text = seg.text

    doc.save(str(path))
    return str(path)
